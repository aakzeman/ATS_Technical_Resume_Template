from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_resume(file_path, content):
    # Create a new Document
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Define formatting styles
    def format_header(run, font_size=Pt(12)):
        run.bold = True
        font = run.font
        font.size = font_size
        font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # Blue color

    def format_body(run):
        run.bold = False
        font = run.font
        font.size = Pt(10)

    def add_section_header(doc, header_text, font_size=Pt(12), alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(header_text)
        format_header(run, font_size)
        # Add underline that spans the entire page
        paragraph._p.get_or_add_pPr().append(parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="4224E9"/></w:pBdr>' % nsdecls('w')))

    def add_section_body(doc, body_text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(body_text)
        format_body(run)

    def add_bullet_points(doc, items):
        for item in items:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(item)
            format_body(run)
            paragraph.style = 'List Bullet'

    # Add sections to the document
    add_section_header(doc, "Name and Contact Information", font_size=Pt(16), alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_section_body(doc, content["contact_info"])

    add_section_header(doc, "Summary")
    add_section_body(doc, content["summary"])

    add_section_header(doc, "Key Core Competencies")
    add_bullet_points(doc, content["competencies"])

    add_section_header(doc, "Skills")
    add_bullet_points(doc, content["skills"])

    add_section_header(doc, "Professional Experience")
    for job in content["experience"]:
        add_section_body(doc, job)

    add_section_header(doc, "Education")
    add_section_body(doc, content["education"])

    # Save the document
    doc.save(file_path)

# Define the content for the resume
content = {
    "contact_info": "EMAIL | PHONE | LOCATION",
    "summary": "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.",
    "competencies": ["COMPETENCY", "COMPETENCY", "COMPETENCY"],
    "skills": ["SKILL", "SKILL", "SKILL"],
    "experience": [
        "TITLE | ORGANIZATION | LOCATION | DATE\nJOB EXPERIENCE\nJOB EXPERIENCE\nJOB EXPERIENCE",
        "TITLE | ORGANIZATION | LOCATION | DATE\nJOB EXPERIENCE\nJOB EXPERIENCE\nJOB EXPERIENCE",
        "TITLE | ORGANIZATION | LOCATION | DATE\nJOB EXPERIENCE\nJOB EXPERIENCE\nJOB EXPERIENCE"
    ],
    "education": "DEGREE | INSTITUTION | DATE"
}

# Create the resume
create_resume("C:/Users/aakzeman/Documents/ATS_Resume_Template_Python.docx", content)
